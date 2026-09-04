"""knowledge_graph 插件 M2 阶段单测 —— 纯逻辑层(不依赖真实 Neo4j/Milvus)。

覆盖:名称归一化 / 结构感知切块 / 文档解析 / 元数据存储 / 本体模板 /
API 层(Fake 图存储 + 临时库)。
真实基础设施的联通与 Cypher 语义由 M2-6 脚本化验收完成(见提交记录)。
"""
import io

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from domains.knowledge_graph.doc_parser import (
    allowed_extension, chunk_text, mime_for, parse_to_text,
)
from domains.knowledge_graph.graph_store import normalize_name
from domains.knowledge_graph.schema_templates import TEMPLATES, get_template_schema, list_templates
from domains.knowledge_graph.store import KGStore


# ── 名称归一化(实体合并锚点) ─────────────────────────────────

class TestNormalizeName:

    def test_basic(self):
        assert normalize_name("  Neo4j ") == "neo4j"
        assert normalize_name("") == ""

    def test_fullwidth(self):
        # 全角字母/数字/空格 → 半角
        assert normalize_name("Ｎｅｏ４ｊ") == "neo4j"
        assert normalize_name("知识　图谱") == "知识 图谱"

    def test_chinese_untouched(self):
        assert normalize_name("知识图谱插件") == "知识图谱插件"

    def test_merge_anchor_stability(self):
        # 同一实体的不同书写应归一到同一锚点
        assert normalize_name("Ｎｅｏ4ｊ") == normalize_name("neo4j ")


# ── 结构感知切块 ─────────────────────────────────────────────

class TestChunkText:

    def test_empty(self):
        assert chunk_text("") == []
        assert chunk_text("   \n\n  ") == []

    def test_heading_boundary(self):
        """标题块 = 软边界:两个章节的内容不粘进同一块。"""
        text = "# 甲\n" + "甲的内容。" * 100 + "\n# 乙\n" + "乙的内容。" * 100
        chunks = chunk_text(text, target_chars=200, overlap_chars=0, max_chars=500)
        assert len(chunks) >= 2
        assert chunks[0]["text"].startswith("# 甲")
        assert any(c["text"].startswith("# 乙") for c in chunks)

    def test_seq_continuous_and_counts(self):
        text = "\n\n".join(f"第{i}段。" + "内容" * 50 for i in range(20))
        chunks = chunk_text(text, target_chars=300, overlap_chars=0, max_chars=600)
        assert [c["seq"] for c in chunks] == list(range(len(chunks)))
        for c in chunks:
            assert c["char_count"] == len(c["text"])

    def test_hard_split_on_oversized_paragraph(self):
        """无结构的超长文本按句读细切,单块不超硬上限。"""
        text = "一句话。" * 2000
        chunks = chunk_text(text, target_chars=400, overlap_chars=0, max_chars=500)
        assert len(chunks) > 1
        assert all(len(c["text"]) <= 500 for c in chunks)
        assert "".join(c["text"] for c in chunks).count("一句话") >= 2000

    def test_overlap(self):
        text = "段落A。" * 300 + "\n\n" + "段落B。" * 300
        chunks = chunk_text(text, target_chars=400, overlap_chars=50, max_chars=800)
        assert len(chunks) >= 2
        # 后一块应包含前一块的尾部片段(重叠生效)
        assert chunks[1]["text"][:10] in chunks[0]["text"] or chunks[1]["text"].startswith(chunks[0]["text"][-50:])

    def test_params_guarded(self):
        """非法参数被夹取,不炸。"""
        chunks = chunk_text("内容。" * 10, target_chars=0, overlap_chars=-5, max_chars=1)
        assert all(len(c["text"]) >= 1 for c in chunks)


# ── 文档解析 ─────────────────────────────────────────────────

class TestParse:

    def test_md_with_bom(self):
        assert "标题" in parse_to_text("a.md", "\ufeff# 标题\n正文".encode("utf-8"))

    def test_gbk_fallback(self):
        assert "中文" in parse_to_text("a.txt", "中文内容".encode("gbk"))

    def test_docx(self):
        import docx as docx_lib
        d = docx_lib.Document()
        d.add_heading("标题一", level=1)
        d.add_paragraph("正文内容")
        buf = io.BytesIO()
        d.save(buf)
        text = parse_to_text("a.docx", buf.getvalue())
        assert "# 标题一" in text and "正文内容" in text

    def test_bad_extension_rejected(self):
        with pytest.raises(ValueError):
            parse_to_text("a.exe", b"MZ")

    def test_whitelist_and_mime(self):
        assert allowed_extension("a.PDF") and allowed_extension("b.md")
        assert not allowed_extension("c.exe") and not allowed_extension("d.doc")
        assert mime_for("x.docx").endswith("wordprocessingml.document")


# ── 元数据存储 ───────────────────────────────────────────────

class TestKGStore:

    @pytest.fixture()
    def store(self, tmp_path):
        return KGStore(str(tmp_path / "kg.db"))

    def test_kb_crud(self, store):
        kb = store.create_kb("库A", "描述", schema_json={"entity_types": [1]}, schema_template="general")
        assert kb["name"] == "库A" and kb["schemaTemplate"] == "general"
        assert store.get_kb(kb["id"])["schema"] == {"entity_types": [1]}
        assert store.get_kb_by_name("库A")["id"] == kb["id"]
        assert store.update_kb(kb["id"], name="库B", description="新描述") is True
        assert store.get_kb(kb["id"])["name"] == "库B"
        items = store.list_kbs()
        assert items[0]["docCount"] == 0 and items[0]["entityTotal"] == 0
        assert store.delete_kb(kb["id"]) is True
        assert store.get_kb(kb["id"]) is None

    def test_document_lifecycle(self, store):
        kb = store.create_kb("库")
        doc = store.create_document(kb["id"], "a.md", "text/markdown", 100, "/tmp/a.md", "hash1")
        assert doc["importStatus"] == "uploaded"
        assert store.get_document_by_hash(kb["id"], "hash1")["id"] == doc["id"]
        store.update_document(doc["id"], import_status="importing", error="x")
        assert store.get_document(doc["id"])["importStatus"] == "importing"
        with pytest.raises(ValueError):
            store.update_document(doc["id"], bad_column=1)  # 列白名单
        store.delete_document(doc["id"])
        assert store.get_document(doc["id"]) is None

    def test_chunks_checkpoint(self, store):
        kb = store.create_kb("库")
        doc = store.create_document(kb["id"], "a.md", "text/markdown", 1, "p", "h")
        n = store.replace_chunks(doc["id"], kb["id"], [
            {"seq": 0, "text": "零"}, {"seq": 1, "text": "一"}, {"seq": 2, "text": "二"},
        ])
        assert n == 3 and store.get_document(doc["id"])["chunkCount"] == 3
        assert [c["seq"] for c in store.list_chunks(doc["id"])] == [0, 1, 2]
        store.mark_chunk(store.list_chunks(doc["id"])[1]["id"], "done")
        assert [c["status"] for c in store.list_chunks(doc["id"])] == ["pending", "done", "pending"]
        assert [c["seq"] for c in store.list_chunks(doc["id"], status="done")] == [1]
        # 重导入路径:replace 重建,chunk_count 同步
        store.replace_chunks(doc["id"], kb["id"], [{"seq": 0, "text": "新"}])
        assert store.get_document(doc["id"])["chunkCount"] == 1


# ── 本体模板 ─────────────────────────────────────────────────

class TestSchemaTemplates:

    def test_templates_registered(self):
        keys = {t["key"] for t in list_templates()}
        assert {"general", "org_people", "product_doc", "regulation"} <= keys

    def test_deep_copy_independence(self):
        s1 = get_template_schema("general")
        s1["entity_types"][0]["label"] = "被篡改"
        s2 = get_template_schema("general")
        assert s2["entity_types"][0]["label"] != "被篡改"

    def test_unknown_falls_back_general(self):
        assert get_template_schema("nope")["entity_types"] == TEMPLATES["general"]["schema"]["entity_types"]


# ── API 层(Fake 图存储 + 临时库) ─────────────────────────────

class FakeGraphStore:
    """内存图存储(只实现 API 层用到的方法)。"""

    def __init__(self):
        self.entities: dict = {}   # kb_id -> {node_id: node}
        self.edges: list = []      # (kb_id, source, target, props)

    def counts(self, kb_id):
        return {"entities": len(self.entities.get(kb_id, {})),
                "relations": sum(1 for e in self.edges if e[0] == kb_id)}

    def get_graph(self, kb_id, q="", node_types=None, limit_nodes=80, limit_edges=150):
        nodes = [n for n in self.entities.get(kb_id, {}).values()
                 if (not q or q in n["name"]) and (not node_types or n["type"] in node_types)]
        ids = {n["id"] for n in nodes}
        edges = [{"id": e[3].get("id", ""), "source": e[1], "target": e[2],
                  "type": e[3].get("type", ""), "description": "", "evidence": "", "docId": ""}
                 for e in self.edges if e[0] == kb_id and e[1] in ids and e[2] in ids]
        return {"nodes": nodes[:limit_nodes], "edges": edges[:limit_edges]}

    def expand_node(self, kb_id, node_id, **kw):
        nodes, edges = [], []
        for e in self.edges:
            if e[0] == kb_id and node_id in (e[1], e[2]):
                other = e[2] if e[1] == node_id else e[1]
                nodes.append(self.entities[kb_id][other])
                edges.append({"id": "", "source": e[1], "target": e[2],
                              "type": e[3].get("type", ""), "description": "", "evidence": "", "docId": ""})
        return {"nodes": nodes, "edges": edges}

    def delete_document(self, kb_id, doc_id):
        self.edges = [e for e in self.edges if e[0] != kb_id or e[3].get("doc_id") != doc_id]
        return {}

    def delete_kb(self, kb_id):
        self.entities.pop(kb_id, None)
        self.edges = [e for e in self.edges if e[0] != kb_id]
        return {}


@pytest.fixture()
def kg_client(tmp_path, monkeypatch):
    """组装最小应用:真实 KGStore(临时库) + Fake 图存储 + 真实 api router。"""
    monkeypatch.setenv("DATABASE_PATH", str(tmp_path / "kg.db"))
    monkeypatch.setenv("KG_FILES_DIR", str(tmp_path / "files"))
    from domains.knowledge_graph import runtime
    runtime.reset_runtime_cache()

    fake = FakeGraphStore()
    monkeypatch.setattr(runtime, "get_graph", lambda state: fake)

    from domains.knowledge_graph.api import router
    app = FastAPI()
    app.include_router(router, prefix="/api/packs/knowledge_graph")
    client = TestClient(app)
    client.app_state_holder = app.state
    yield client, fake
    runtime.reset_runtime_cache()


class TestKnowledgeGraphApi:

    def test_kb_crud_flow(self, kg_client):
        client, fake = kg_client
        r = client.post("/api/packs/knowledge_graph/kbs",
                        json={"name": "测试库", "template": "product_doc"})
        assert r.status_code == 200
        kb = r.json()
        assert len(kb["schema"]["entity_types"]) == len(TEMPLATES["product_doc"]["schema"]["entity_types"])
        # graph 计数在详情端点(GET /kbs/{id})
        detail = client.get(f"/api/packs/knowledge_graph/kbs/{kb['id']}").json()
        assert detail["graph"] == {"entities": 0, "relations": 0}

        assert client.post("/api/packs/knowledge_graph/kbs",
                           json={"name": "测试库"}).status_code == 409
        assert client.post("/api/packs/knowledge_graph/kbs",
                           json={"name": " "}).status_code == 422
        assert client.get(f"/api/packs/knowledge_graph/kbs/不存在").status_code == 404

        r = client.put(f"/api/packs/knowledge_graph/kbs/{kb['id']}",
                       json={"description": "改", "schema": {"schema_mode": "strict"}})
        assert r.json()["description"] == "改" and r.json()["schema"]["schema_mode"] == "strict"

        assert client.delete(f"/api/packs/knowledge_graph/kbs/{kb['id']}").json()["success"] is True

    def test_upload_dedup_and_reject(self, kg_client):
        client, _ = kg_client
        kb = client.post("/api/packs/knowledge_graph/kbs", json={"name": "U"}).json()
        base = "/api/packs/knowledge_graph/kbs"
        content = "# A\n内容".encode()
        r1 = client.post(f"{base}/{kb['id']}/documents",
                         files=[("files", ("a.md", content, "text/markdown"))]).json()["items"]
        assert r1[0]["ok"] and r1[0]["document"]["contentHash"]
        r2 = client.post(f"{base}/{kb['id']}/documents",
                         files=[("files", ("b.md", content, "text/markdown"))]).json()["items"]
        assert not r2[0]["ok"] and "已存在" in r2[0]["reason"]
        r3 = client.post(f"{base}/{kb['id']}/documents",
                         files=[("files", ("c.exe", b"MZ", "application/octet-stream"))]).json()["items"]
        assert not r3[0]["ok"] and "格式" in r3[0]["reason"]

    def test_graph_view_and_expand(self, kg_client):
        client, fake = kg_client
        kb = client.post("/api/packs/knowledge_graph/kbs", json={"name": "G"}).json()
        fake.entities[kb["id"]] = {
            "n1": {"id": "n1", "name": "Alpha", "type": "component",
                   "description": "", "aliases": [], "sourceDocs": [], "typeStatus": "", "updatedAt": ""},
            "n2": {"id": "n2", "name": "Beta", "type": "product",
                   "description": "", "aliases": [], "sourceDocs": [], "typeStatus": "", "updatedAt": ""},
        }
        fake.edges.append((kb["id"], "n1", "n2", {"type": "依赖", "doc_id": "d1"}))

        g = client.get(f"/api/packs/knowledge_graph/kbs/{kb['id']}/graph").json()
        assert len(g["nodes"]) == 2 and len(g["edges"]) == 1
        g = client.get(f"/api/packs/knowledge_graph/kbs/{kb['id']}/graph?q=Alph").json()
        assert [n["id"] for n in g["nodes"]] == ["n1"] and g["edges"] == []
        g = client.get(f"/api/packs/knowledge_graph/kbs/{kb['id']}/graph?types=product").json()
        assert [n["id"] for n in g["nodes"]] == ["n2"]
        ex = client.get(f"/api/packs/knowledge_graph/kbs/{kb['id']}/graph/expand?node_id=n1").json()
        assert [n["id"] for n in ex["nodes"]] == ["n2"]
        assert client.get(f"/api/packs/knowledge_graph/kbs/{kb['id']}/graph/expand").status_code == 422
