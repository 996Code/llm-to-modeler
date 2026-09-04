"""文档解析与结构感知切块(导入流水线的第 1-2 步)。

【支持格式】
  .md / .markdown / .txt  → UTF-8 直读(容错 BOM / GBK 兜底解码)
  .pdf                    → pypdf 逐页抽取
  .docx                   → python-docx 段落抽取(表格单元格拼接)

【切块策略】(结构感知,目标 target 字数)
  1. md 按标题行(#/##/...)切段落块,纯文本按空行分段
  2. 段落贪心装箱到 target;超过 max 硬上限的段落再按句号/换行细切
  3. 相邻块之间保留 overlap 字数的尾部(跨块语义连续)
  4. 每块带 seq(0 起),供 chunk 级 checkpoint 引用
"""
import io
import logging
import re
from typing import Any, Dict, List

logger = logging.getLogger(__name__)

# 扩展名白名单(上传校验与解析分发共用)
ALLOWED_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf", ".docx"}

MIME_BY_EXT = {
    ".md": "text/markdown", ".markdown": "text/markdown",
    ".txt": "text/plain", ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_HEADING_RE = re.compile(r"^(#{1,6})\s+")


def allowed_extension(filename: str) -> bool:
    name = (filename or "").lower()
    return any(name.endswith(ext) for ext in ALLOWED_EXTENSIONS)


def mime_for(filename: str) -> str:
    name = (filename or "").lower()
    for ext, mime in MIME_BY_EXT.items():
        if name.endswith(ext):
            return mime
    return "application/octet-stream"


# ── 解析:文件字节 → 纯文本 ─────────────────────────────────

def parse_to_text(filename: str, data: bytes) -> str:
    """按扩展名解析文件为纯文本(无法解析抛 ValueError,由任务层记失败)。"""
    name = (filename or "").lower()
    if name.endswith((".md", ".markdown", ".txt")):
        return _decode_text(data)
    if name.endswith(".pdf"):
        return _parse_pdf(data)
    if name.endswith(".docx"):
        return _parse_docx(data)
    raise ValueError(f"不支持的文件类型: {filename}(支持 md/txt/pdf/docx)")


def _decode_text(data: bytes) -> str:
    """UTF-8 优先,失败回退 GBK(中文办公环境常见),再失败用 errors=replace。

    BOM/零字节探测:UTF-16/UTF-32 字节流常能"成功"通过 GBK 解码成含 NUL
    的乱码(不抛异常),这类垃圾进切块再进 LLM 就是一张废图谱——所以
    GBK 解出来后若含 NUL 或高占比替换符,按解码失败处理。
    """
    if data[:3] == b"\xef\xbb\xbf":
        return data[3:].decode("utf-8", errors="replace")
    if data[:2] in (b"\xff\xfe", b"\xfe\xff"):
        return data.decode("utf-16", errors="replace")
    if data[:4] in (b"\xff\xfe\x00\x00", b"\x00\x00\xfe\xff"):
        return data.decode("utf-32", errors="replace")
    if b"\x00" in data[:200]:
        # 无 BOM 但含零字节:未知宽字符编码,按 UTF-16 兜底再校验
        guess = data.decode("utf-16", errors="ignore")
        if guess and "\x00" not in guess:
            return guess
        raise ValueError("无法识别的文本编码(疑似宽字符/二进制内容)")
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = data.decode("gbk")
            if "\x00" in text:
                raise ValueError("无法识别的文本编码(疑似宽字符/二进制内容)")
            return text
        except UnicodeDecodeError:
            fallback = data.decode("utf-8", errors="replace")
            # 替换符占比过高 = 实际上没解出来,别把乱码当文本
            if fallback.count("\ufffd") > len(fallback) * 0.05:
                raise ValueError("无法识别的文本编码")
            return fallback


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # 单页损坏不整体失败
            logger.warning(f"pdf 第 {i + 1} 页解析失败(跳过): {e}")
            text = ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


def _parse_docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts: List[str] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if text:
            style = (para.style.name or "").lower()
            # 保留标题层级信息(切块的标题边界用)
            if "heading 1" in style:
                parts.append("# " + text)
            elif "heading 2" in style:
                parts.append("## " + text)
            elif "heading 3" in style:
                parts.append("### " + text)
            else:
                parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n\n".join(parts)


# ── 切块:纯文本 → chunks ────────────────────────────────────

def _split_blocks(text: str) -> List[str]:
    """切成"结构块":标题行独立成块(保留 # 前缀),其余按空行分段。"""
    blocks: List[str] = []
    current: List[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if _HEADING_RE.match(stripped):
            if current:
                blocks.append("\n".join(current).strip())
                current = []
            blocks.append(stripped)  # 标题独立成块(软边界)
        elif stripped == "":
            if current:
                blocks.append("\n".join(current).strip())
                current = []
        else:
            current.append(line)
    if current:
        blocks.append("\n".join(current).strip())
    return [b for b in blocks if b]


def _hard_split(block: str, max_chars: int) -> List[str]:
    """超长块按句读边界细切(句号/问叹号/分号/换行),仍超长则暴力截断。"""
    if len(block) <= max_chars:
        return [block]
    sentences = re.split(r"(?<=[。！？；;!?])\s*|\n", block)
    parts: List[str] = []
    buf = ""
    for sent in sentences:
        if not sent:
            continue
        while len(sent) > max_chars:  # 单句超长:暴力切
            if buf:
                parts.append(buf)
                buf = ""
            parts.append(sent[:max_chars])
            sent = sent[max_chars:]
        if len(buf) + len(sent) > max_chars and buf:
            parts.append(buf)
            buf = sent
        else:
            buf = buf + ("\n" if buf else "") + sent if buf else sent
    if buf:
        parts.append(buf)
    return parts


def chunk_text(
    text: str,
    target_chars: int = 1200,
    overlap_chars: int = 100,
    max_chars: int = 3000,
) -> List[Dict[str, Any]]:
    """结构感知切块。Returns: [{seq, text, char_count}](seq 从 0 连续)。"""
    if not text or not text.strip():
        return []

    # 归一参数(防呆:target/overlap/max 的非法配置不至于炸掉导入)
    target_chars = max(200, int(target_chars or 1200))
    max_chars = max(target_chars, int(max_chars or 3000))
    overlap_chars = max(0, min(int(overlap_chars or 0), target_chars // 2))

    chunks: List[str] = []
    buf = ""
    for block in _split_blocks(text):
        # 标题块 = 软边界:直接结算当前缓冲(避免跨标题粘连)
        if _HEADING_RE.match(block) and buf:
            chunks.append(buf)
            buf = ""
        for piece in _hard_split(block, max_chars):
            if len(buf) + len(piece) + 1 > target_chars and buf:
                chunks.append(buf)
                buf = piece
            else:
                buf = buf + "\n" + piece if buf else piece
    if buf:
        chunks.append(buf)

    # 相邻块重叠:后块带上前块尾部(跨块语义连续;标题块不加,避免污染)
    if overlap_chars > 0:
        overlapped: List[str] = []
        for i, c in enumerate(chunks):
            if i == 0 or _HEADING_RE.match(c):
                overlapped.append(c)
                continue
            prev_tail = chunks[i - 1][-overlap_chars:]
            overlapped.append(prev_tail + "\n" + c)
        chunks = overlapped

    # 硬上限最终兜底(重叠可能超)
    final: List[str] = []
    for c in chunks:
        final.extend(_hard_split(c, max_chars))

    return [
        {"seq": i, "text": t, "char_count": len(t)}
        for i, t in enumerate(final) if t.strip()
    ]
